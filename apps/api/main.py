"""Hiro2 API（Phase A，ADR 0006）：只读 View Model + append-only 审核。

启动：
    uv run uvicorn apps.api.main:app --port 8000
前端联调（apps/web）：
    NEXT_PUBLIC_USE_MOCK=false NEXT_PUBLIC_API_BASE_URL=http://localhost:8000
"""

from __future__ import annotations

import asyncio
import os
from contextlib import asynccontextmanager
from pathlib import Path

from fastapi import FastAPI, HTTPException, Response, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse, PlainTextResponse
from pydantic import BaseModel, Field

from backend.application.annotate import submit_annotation
from backend.application.career import (
    add_proof,
    get_career_home,
    save_growth_task,
    save_profile,
    set_active_target,
)
from backend.application.dashboard import build_dashboard
from backend.application.datasets import build_dataset_overview, build_dataset_overview_db
from backend.application.diagnosis import build_diagnosis, list_candidates
from backend.application.evaluation import build_evaluation_overview
from backend.application.insights import build_detected_changes, build_timeline
from backend.application.joblist import build_published_jobs
from backend.application.pipeline_runs import build_pipeline_runs
from backend.application.quality import build_quality_overview
from backend.application.service import ApplicationService
from backend.application.outbox_worker import outbox_worker_enabled, outbox_worker_loop
from backend.application.snapshot import snapshot_enabled, snapshot_loop
from backend.application.temporal_vm import build_skill_graph, build_tasks, build_temporal
from backend.application.training import build_training_output


@asynccontextmanager
async def _lifespan(_: FastAPI):
    tasks: list[asyncio.Task] = []
    # 配置了 PG 时启动即后台跑一次数据导入（幂等，不阻塞就绪）
    if os.getenv("DATABASE_URL"):
        from backend.application.snapshot import run_import_once

        tasks.append(asyncio.create_task(run_import_once()))
        print("[import] 后台数据导入任务已启动", flush=True)
    # JD 快照后台任务：HIRO2_SNAPSHOT_ENABLED=true 时启动即采集（周期见模块头）
    if snapshot_enabled():
        tasks.append(asyncio.create_task(snapshot_loop()))
        print("[snapshot] 后台采集任务已启动", flush=True)
    # Outbox worker：发布事件 -> Neo4j 投影自动消费（带退避重试）
    if os.getenv("DATABASE_URL") and outbox_worker_enabled():
        tasks.append(asyncio.create_task(outbox_worker_loop()))
        print("[outbox] 后台消费 worker 已启动", flush=True)
    yield
    for t in tasks:
        t.cancel()


app = FastAPI(title="Hiro2 API", version="0.1.0", lifespan=_lifespan)


def _cors_origins() -> list[str]:
    configured = os.getenv("CORS_ORIGINS", "")
    extra = [origin.strip().rstrip("/") for origin in configured.split(",") if origin.strip()]
    return ["http://localhost:3000", "http://127.0.0.1:3000", *extra]


app.add_middleware(
    CORSMiddleware,
    allow_origins=_cors_origins(),
    allow_origin_regex=os.getenv(
        "CORS_ORIGIN_REGEX",
        r"https?://10\.\d+\.\d+\.\d+(?::\d+)?$",
    ),
    allow_methods=["*"],
    allow_headers=["*"],
)
svc = ApplicationService()


class ReviewRequest(BaseModel):
    decision: str = Field(pattern="^(accepted|rejected|needs_evidence|modified)$")
    note: str = ""


class PublishRequest(BaseModel):
    reviewer: str = ""
    note: str = ""


class GrowthTaskRequest(BaseModel):
    completed: bool


class ProofRequest(BaseModel):
    skill_id: str = Field(min_length=1, max_length=80)
    title: str = Field(min_length=1, max_length=120)
    description: str = Field(default="", max_length=1000)
    proof_url: str | None = Field(default=None, max_length=500)


class TargetRequest(BaseModel):
    job_version_id: str = Field(min_length=1, max_length=120)


class ProfileSkillRequest(BaseModel):
    name: str = Field(min_length=1, max_length=80)
    status: str = Field(pattern="^(ready|partial|missing)$")


class ProfileRequest(BaseModel):
    skills: list[ProfileSkillRequest] = Field(default_factory=list, max_length=60)
    projects: list[str] = Field(default_factory=list, max_length=10)


def _health_status() -> dict:
    postgres = "not_configured"
    database_url = os.getenv("DATABASE_URL", "")
    if database_url:
        try:
            import psycopg

            with psycopg.connect(database_url, connect_timeout=2):
                postgres = "ok"
        except Exception:
            postgres = "unavailable"

    neo4j = "not_configured"
    if os.getenv("NEO4J_URI"):
        try:
            from backend.infra.neo4j import check_neo4j

            neo4j = "ok" if check_neo4j() else "unavailable"
        except Exception:
            neo4j = "unavailable"

    required = [postgres] if database_url else []
    status = "ok" if all(x == "ok" for x in required) else "degraded"
    return {"status": status, "postgres": postgres, "neo4j": neo4j}


@app.get("/health/live")
def health_live() -> dict:
    """Process liveness; does not require external dependencies."""
    return {"status": "ok"}


@app.get("/health/ready")
def health_ready(response: Response) -> dict:
    """Readiness; PostgreSQL is required when configured, Neo4j is degradable."""
    result = _health_status()
    if result["status"] != "ok":
        response.status_code = 503
    return result


@app.get("/health")
def health(response: Response) -> dict:
    """Backward-compatible alias for readiness checks."""
    return health_ready(response)


@app.get("/api/v1/jobs/default/update")
def job_update(state: str = "ready") -> dict:
    if state not in ("ready", "empty", "error"):
        raise HTTPException(400, "state 只支持 ready/empty/error")
    if state == "error":
        raise HTTPException(503, "模拟错误态（前端 error 边界联调）")
    return svc.job_update(state).model_dump(by_alias=True)


@app.get("/api/v1/emerging-jobs")
def emerging_jobs() -> dict:
    return svc.emerging_jobs().model_dump(by_alias=True)


@app.get("/api/v1/dashboard/overview")
def dashboard_overview() -> dict:
    return build_dashboard().model_dump()


@app.get("/api/v1/jobs/detected-changes")
def jobs_detected_changes() -> dict:
    """快照 Diff 自动检测的岗位变化草稿（PENDING，审核后喂 jobver 升版）。"""
    return build_detected_changes().model_dump()


@app.get("/api/v1/jobs/published")
def jobs_published() -> dict:
    """已发布岗位版本列表（每岗位取最新版本），求职区目标岗位页数据源。"""
    return build_published_jobs().model_dump()


@app.get("/api/v1/temporal/timeline")
def temporal_timeline() -> dict:
    """四层时间轴：论文 arXiv -> PyPI/npm 包 -> 日报 -> JD 的传导。"""
    return build_timeline().model_dump()


@app.get("/api/v1/datasets/overview")
def datasets_overview() -> dict:
    dsn = os.getenv("DATABASE_URL")
    if dsn:
        try:
            return build_dataset_overview_db(dsn).model_dump()
        except Exception:
            pass
    return build_dataset_overview().model_dump()


@app.get("/api/v1/pipeline-runs")
def pipeline_runs(limit: int = 50, since_days: int = 7) -> dict:
    """最近 pipeline run 列表（默认 7 天 / 50 条）。只读，扫描 data/runs/。"""
    return build_pipeline_runs(limit=limit, since_days=since_days).model_dump()


@app.get("/api/v1/evaluation/overview")
def evaluation_overview() -> dict:
    return build_evaluation_overview()


@app.get("/api/v1/emerging-jobs/{candidate_id}/review")
def emerging_review_get(candidate_id: str) -> dict:
    data = svc.emerging_jobs().model_dump(by_alias=True)
    for c in data["candidates"]:
        if c["id"] == candidate_id:
            return c
    raise HTTPException(404, f"候选不存在: {candidate_id}")


@app.post("/api/v1/emerging-jobs/{candidate_id}/review")
def emerging_review(candidate_id: str, req: ReviewRequest) -> dict:
    return svc.submit_review(candidate_id, req.decision, req.note)


@app.get("/api/v1/evidence/{evidence_id}")
def evidence(evidence_id: str) -> dict:
    vm = svc.evidence_by_id(evidence_id)
    if vm is None:
        raise HTTPException(404, f"证据不存在: {evidence_id}")
    return vm.model_dump(by_alias=True)


@app.get("/api/v1/career/home")
def career_home() -> dict:
    return get_career_home()


@app.get("/api/v1/candidates")
def candidates() -> list[dict]:
    return list_candidates()


@app.get("/api/v1/diagnosis/{candidate_id}")
def diagnosis(candidate_id: str, job: str = "ai-agent-v2") -> dict:
    try:
        return build_diagnosis(candidate_id, job).model_dump(by_alias=True)
    except FileNotFoundError as exc:
        raise HTTPException(404, str(exc)) from exc


@app.put("/api/v1/candidates/{candidate_id}/growth-tasks/{skill_id}")
def update_growth_task(
    candidate_id: str, skill_id: str, req: GrowthTaskRequest, job: str = "ai-agent-v2"
) -> dict:
    try:
        return save_growth_task(candidate_id, job, skill_id, req.completed)
    except RuntimeError as exc:
        raise HTTPException(503, str(exc)) from exc


@app.post("/api/v1/candidates/{candidate_id}/proofs")
def create_proof(candidate_id: str, req: ProofRequest) -> dict:
    try:
        return add_proof(candidate_id, req.skill_id, req.title, req.description, req.proof_url)
    except RuntimeError as exc:
        raise HTTPException(503, str(exc)) from exc


@app.put("/api/v1/candidates/{candidate_id}/target")
def update_target(candidate_id: str, req: TargetRequest) -> dict:
    try:
        return set_active_target(candidate_id, req.job_version_id)
    except RuntimeError as exc:
        raise HTTPException(503, str(exc)) from exc


@app.patch("/api/v1/candidates/{candidate_id}/profile")
def update_profile(candidate_id: str, req: ProfileRequest) -> dict:
    try:
        return save_profile(
            candidate_id,
            [skill.model_dump() for skill in req.skills],
            req.projects,
        )
    except LookupError as exc:
        raise HTTPException(404, str(exc)) from exc
    except RuntimeError as exc:
        raise HTTPException(503, str(exc)) from exc


@app.get("/api/v1/temporal/dataset")
def temporal_dataset() -> dict:
    data = build_temporal().model_dump()
    # 前端契约：顶层 camelCase，内层字段保持 snake_case
    data["backtestRecords"] = data.pop("backtest_records")
    return data


class SuggestionReviewRequest(BaseModel):
    decision: str = Field(pattern="^(accepted|rejected|needs_evidence|modified)$")
    note: str = ""
    suggested_level: str = ""


@app.post("/api/v1/temporal/suggestions/{suggestion_id}/review")
def temporal_suggestion_review(suggestion_id: str, req: SuggestionReviewRequest) -> dict:
    """岗位影响建议审核：写入 append-only 审核日志（与岗位审核同一份事实）。

    边界：建议被接受不直接修改岗位版本，只标记建议状态；
    后续岗位变更仍须走招聘区的 Diff 审核与发布流程。
    """
    known = {s.suggestion_id for s in build_temporal().suggestions}
    if suggestion_id not in known:
        raise HTTPException(404, f"建议不存在: {suggestion_id}")
    extra = {"suggested_level": req.suggested_level} if req.suggested_level else {}
    action = svc.submit_review(suggestion_id, req.decision, req.note, **extra)
    return {**action, "suggestion_id": suggestion_id}


@app.get("/api/v1/skills/graph")
def skills_graph(job: str = "ai-agent-v2") -> dict:
    graph = build_skill_graph(job)
    if os.getenv("NEO4J_URI"):
        try:
            from backend.infra.neo4j import read_job_graph

            nodes, _ = read_job_graph(job)
            if nodes:
                # Neo4j 投影校准 role；VM 节点的别名/位置/关联版本/信号等字段保留
                roles = {n["id"]: n["role"] for n in nodes if n["id"] != "root"}
                graph = graph.model_copy(
                    update={
                        "nodes": [
                            n.model_copy(update={"role": roles[n.id]}) if n.id in roles else n
                            for n in graph.nodes
                        ]
                    }
                )
        except Exception:
            pass
    raw = graph.model_dump()
    return {
        "fixtureVersion": raw["fixture_version"],
        "mode": raw["mode"],
        "run": {
            "id": raw["run"].get("id", ""),
            "datasetVersion": raw["run"].get("datasetVersion", ""),
            "status": raw["run"].get("status", "REVIEWING"),
        },
        "context": {
            "jobTitle": raw["context"].get("jobTitle", ""),
            "baselineVersion": raw["context"].get("baselineVersion", ""),
            "targetVersion": raw["context"].get("targetVersion", ""),
            "timeWindow": raw["context"].get("timeWindow", ""),
        },
        "nodes": [
            {
                "id": node["id"],
                "label": node["label"],
                "capabilityId": node["capability_id"],
                "pointName": node["point_name"],
                "role": node["role"],
                "status": node["status"],
                "aliases": node["aliases"],
                "evidenceIds": node["evidence_ids"],
                "position": node["position"],
                "techStack": node["tech_stack"],
                "jobVersions": [
                    {
                        "versionId": ref["version_id"],
                        "title": ref["title"],
                        "role": ref["role"],
                        "weight": ref["weight"],
                    }
                    for ref in node["job_versions"]
                ],
                "signal": {
                    "jdMentions": node["signal"].get("jd_mentions", 0),
                    "mentionShare": node["signal"].get("mention_share", 0),
                }
                if node["signal"]
                else None,
            }
            for node in raw["nodes"]
        ],
        "edges": raw["edges"],
        "filterOptions": {
            "techStacks": raw["filter_options"].get("techStacks", []),
            "roles": raw["filter_options"].get("roles", []),
            "capabilityTypes": raw["filter_options"].get("capabilityTypes", []),
        },
    }


class TaskDecisionRequest(BaseModel):
    decision: str = Field(pattern="^(ACCEPT|MODIFY|REJECT|UNKNOWN)$")
    rationale: str = ""
    reviewer_id: str = "local"
    error_type: str | None = None
    corrected_payload: dict | None = None


@app.get("/api/v1/tasks/my")
def my_tasks() -> dict:
    return build_tasks().model_dump()


@app.post("/api/v1/tasks/{task_id}/decision")
def task_decision(task_id: str, req: TaskDecisionRequest) -> dict:
    known = {t.task_id for t in build_tasks().tasks}
    if task_id not in known:
        raise HTTPException(404, f"任务不存在: {task_id}")
    rec = submit_annotation(
        task_id,
        req.decision,
        rationale=req.rationale,
        reviewer_id=req.reviewer_id,
        error_type=req.error_type,
        corrected_payload=req.corrected_payload,
    )
    return {
        "task_id": task_id,
        "status": "RESOLVED",
        "annotation_id": rec["annotation_id"],
    }


@app.get("/api/v1/quality/overview")
def quality_overview() -> dict:
    return build_quality_overview(os.getenv("DATABASE_URL") or None).model_dump()


@app.get("/api/v1/jobs/{job_version_id}/training-output")
def training_output(job_version_id: str) -> dict:
    """F-T3.6：岗位标准下游输出——JD 模板 / 培养任务 / 能力证明要求。"""
    try:
        return build_training_output(job_version_id).model_dump()
    except FileNotFoundError as exc:
        raise HTTPException(404, str(exc)) from exc


@app.post("/api/v1/jobs/{job_version_id}/versions/{version}/publish")
def publish_job(job_version_id: str, version: str, req: PublishRequest) -> dict:
    """发布流：审核留痕 + jobpub 固化为不可变 PUBLISHED；重复发布幂等返回。

    当前只有 default（AI Agent 主案例草稿）一条发布路径，与
    GET /jobs/default/update 同源；version 路径参数保留给未来多草稿。
    """
    import json
    import re
    import sys

    root = Path(__file__).resolve().parents[2]
    if str(root) not in sys.path:
        sys.path.insert(0, str(root))
    from scripts.jobpub import cmd_publish

    draft_path = root / "data/processed/jd-opencli/jobversion-agent-draft.json"
    if job_version_id != "default" or not draft_path.is_file():
        raise HTTPException(404, f"草稿不存在: {job_version_id}")
    draft = json.loads(draft_path.read_text(encoding="utf-8"))
    vid = re.sub(r"-draft-\d{8}$", "", draft.get("version_id") or version)
    reviewer = req.reviewer or os.getenv("HIRO2_REVIEWER", "webui")
    svc.submit_review(
        draft.get("job_id", job_version_id), "accepted", req.note or "Web 发布流审核通过"
    )

    pub_path = root / "data/processed/jobversions/published" / f"{vid}.json"
    try:
        cmd_publish(draft_path, reviewer, req.note, vid)
    except SystemExit as exc:
        # 已发布过（发布后不可变）：读回幂等返回，而非报错打断 UI 流
        if not pub_path.is_file():
            raise HTTPException(409, str(exc)) from exc
    pub = json.loads(pub_path.read_text(encoding="utf-8"))
    return {
        "versionId": pub["version_id"],
        "publishedAt": pub.get("published_at", ""),
        "reviewActionIds": pub.get("review_action_ids", []),
    }


@app.post("/api/v1/candidates/resumes")
async def upload_resume(file: UploadFile) -> dict:
    """F-T3.4：上传 PDF/DOCX -> LLM 抽取 -> 双层归一 -> 候选人画像。"""
    import tempfile

    from backend.candidates.parse import (
        ResumeRawExtraction,
        build_profile,
        llm_resolve_unmatched,
        parse_document,
    )
    from backend.skills.resolver import load_resolver

    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in (".pdf", ".docx", ".txt", ".md"):
        raise HTTPException(400, f"不支持的格式 {suffix}，请上传 PDF/DOCX/TXT")

    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        content = await file.read()
        tmp.write(content)
        tmp_path = Path(tmp.name)

    try:
        text = parse_document(tmp_path)
        if not text.strip():
            raise HTTPException(422, "解析结果为空，请检查文件内容")
    except ValueError as exc:
        raise HTTPException(400, str(exc))

    # LLM 抽取 + 双层归一（endpoint 本身是 async，直接 await）
    from scripts.candmatch import _extract

    raw = await _extract(text)
    extraction = ResumeRawExtraction.model_validate(raw)
    resolver = load_resolver()
    profile = build_profile("upload_pending", extraction, resolver)

    # LLM 归层
    unresolved = [s.mention for s in profile.skills if not s.skill_id]
    cands: dict = {}
    if unresolved:
        cands = await llm_resolve_unmatched(unresolved, text[:400])
        for s in profile.skills:
            if s.skill_id or s.mention not in cands:
                continue
            c = cands[s.mention]
            if c.is_skill and c.capability_id and c.confidence >= 0.6:
                s.skill_id = c.capability_id
                s.resolved_by = "llm"
                s.reason = c.reason
            else:
                s.resolved_by = "unmatched"

    result = {
        "rawText": text[:2000],
        "profile": profile.model_dump(),
        "stats": {
            "totalSkills": len(profile.skills),
            "resolved": sum(1 for s in profile.skills if s.skill_id),
            "byDict": sum(1 for s in profile.skills if s.resolved_by == "dict"),
            "byLlm": sum(1 for s in profile.skills if s.resolved_by == "llm"),
            "unresolved": sum(1 for s in profile.skills if s.resolved_by == "unmatched"),
        },
    }
    # 入档：文件落 objects/resumes，元数据与画像追加 JSONL（幂等由 resume_id 唯一保证）
    from backend.candidates.archive import save_to_archive

    record = save_to_archive(content, file.filename or "resume", result)
    return {**result, "resumeId": record["resume_id"], "source": "upload"}


@app.get("/api/v1/candidates/resumes")
def list_resume_archive() -> list[dict]:
    """简历档案列表（轻字段，最新在前）。"""
    from backend.candidates.archive import list_archive

    return list_archive()


@app.get("/api/v1/candidates/resumes/{resume_id}")
def get_resume_archive(resume_id: str) -> dict:
    """单个档案详情（含画像与原文片段）；imported 档案未解析时画像为空。"""
    from backend.candidates.archive import get_archive

    record = get_archive(resume_id)
    if record is None:
        raise HTTPException(404, f"档案不存在: {resume_id}")
    return {
        "resumeId": record["resume_id"],
        "filename": record["filename"],
        "size": record["size"],
        "uploadedAt": record["uploaded_at"],
        "source": record["source"],
        "rawText": record.get("raw_text", ""),
        "profile": record.get("profile") or {},
        "stats": record.get("stats"),
    }


@app.get("/api/v1/candidates/resumes/{resume_id}/preview")
def preview_resume_archive(resume_id: str):
    """PDF 原件内嵌；DOCX 转安全只读 HTML；文本文件直接预览。"""
    from html import escape

    from backend.candidates.archive import get_stored_file

    path = get_stored_file(resume_id)
    if path is None:
        raise HTTPException(404, f"原始文件不存在: {resume_id}")
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return FileResponse(path, media_type="application/pdf", content_disposition_type="inline")
    if suffix == ".docx":
        from docx import Document

        doc = Document(str(path))
        blocks = [
            f"<p>{escape(paragraph.text)}</p>"
            for paragraph in doc.paragraphs
            if paragraph.text.strip()
        ]
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = "".join(f"<td>{escape(cell.text)}</td>" for cell in row.cells)
                rows.append(f"<tr>{cells}</tr>")
            blocks.append(f"<table>{''.join(rows)}</table>")
        html = """<!doctype html><meta charset='utf-8'><style>
        body{margin:0;padding:32px;font:14px/1.7 system-ui;color:#20201d;background:#fff}
        p{margin:0 0 10px;white-space:pre-wrap}
        table{border-collapse:collapse;width:100%;margin:12px 0}
        td{border:1px solid #ddd;padding:6px}</style>""" + "".join(blocks)
        return HTMLResponse(html)
    if suffix in (".txt", ".md"):
        return PlainTextResponse(path.read_text(encoding="utf-8"))
    raise HTTPException(415, f"暂不支持预览: {suffix}")
