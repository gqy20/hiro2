"""候选人文档解析与画像构建（薄适配器：PyMuPDF / python-docx）。"""

from __future__ import annotations

from pathlib import Path
from typing import Protocol

from pydantic import BaseModel, ConfigDict, Field


class ResumeRawExtraction(BaseModel):
    """对应 prompts/resume-parse.yml 的 output_schema。"""

    model_config = ConfigDict(extra="forbid")

    skills: list[dict] = Field(default_factory=list, max_length=40)
    experience_years: float | None = None
    education: str = ""
    location: str = ""
    work_experiences: list[dict] = Field(default_factory=list, max_length=10)
    education_history: list[dict] = Field(default_factory=list, max_length=8)
    certificates: list[dict] = Field(default_factory=list, max_length=12)
    portfolio_urls: list[str] = Field(default_factory=list, max_length=8)
    languages: list[str] = Field(default_factory=list, max_length=8)
    projects: list[dict] = Field(default_factory=list, max_length=8)


class DocumentParser(Protocol):
    def parse(self, path: Path) -> str: ...


class PdfParser:
    """PyMuPDF 文本抽取。已知局限：双栏版式阅读顺序可能错乱（升级项 MinerU）。"""

    def parse(self, path: Path) -> str:
        import fitz

        doc = fitz.open(path)
        try:
            return "\n".join(doc[i].get_text("text") for i in range(doc.page_count))
        finally:
            doc.close()


class DocxParser:
    """python-docx：段落 + 表格（DOCX 为流式逻辑格式，天然保序）。"""

    def parse(self, path: Path) -> str:
        from docx import Document

        doc = Document(str(path))
        paragraphs = list(doc.paragraphs)  # 类型桩不完整，显式转 list
        parts = [p.text for p in paragraphs if p.text.strip()]
        tables = list(doc.tables)
        for table in tables:
            for row in list(table.rows):
                cells = [c.text.strip() for c in list(row.cells) if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)


def parse_document(path: Path) -> str:
    """按扩展名分发到对应解析器（薄适配器，AGENTS 规则）。"""
    if path.suffix.lower() == ".pdf":
        return PdfParser().parse(path)
    if path.suffix.lower() == ".docx":
        return DocxParser().parse(path)
    if path.suffix.lower() in (".txt", ".md"):
        return path.read_text(encoding="utf-8")
    raise ValueError(f"不支持的格式: {path.suffix}（.doc 老格式请转存 .docx）")


def build_profile(
    candidate_id: str,
    raw: ResumeRawExtraction,
    resolver,
    corrections: list[dict] | None = None,
) -> CandidateProfile:
    """raw extraction + 用户修正 -> effective profile（原始永不覆盖，可审计）。"""
    from .models import (
        CandidateProfile,
        CertificateEntry,
        EducationEntry,
        EffectiveSkill,
        ProjectEntry,
        WorkExperience,
    )

    skills: list[EffectiveSkill] = []
    seen: set[str] = set()
    for s in raw.skills:
        m = s.get("mention", "")
        if not m or m in seen:
            continue
        seen.add(m)
        hit = resolver.resolve(m)
        skills.append(
            EffectiveSkill(
                mention=m,
                skill_id=hit.skill_id,
                point_id=hit.point_id,
                proficiency=s.get("proficiency", "初级"),
                years=s.get("years"),
                source="raw",
            )
        )
    for c in corrections or []:
        m = c.get("mention", "")
        if not m:
            continue
        hit = resolver.resolve(m)
        skills.append(
            EffectiveSkill(
                mention=m,
                skill_id=hit.skill_id,
                point_id=hit.point_id,
                proficiency=c.get("proficiency", "初级"),
                years=c.get("years"),
                source="correction",
            )
        )
    return CandidateProfile(
        candidate_id=candidate_id,
        raw_extraction_id=f"{candidate_id}:raw",
        skills=skills,
        experience_years=raw.experience_years,
        education=raw.education,
        location=raw.location,
        work_experiences=[WorkExperience.model_validate(item) for item in raw.work_experiences],
        education_history=[EducationEntry.model_validate(item) for item in raw.education_history],
        certificates=[CertificateEntry.model_validate(item) for item in raw.certificates],
        portfolio_urls=raw.portfolio_urls,
        languages=raw.languages,
        projects=[ProjectEntry.model_validate(p) for p in raw.projects],
        correction_log=corrections or [],
    )


from .models import CandidateProfile  # noqa: E402,F401  (re-export for callers)


class ResumeAliasResult(BaseModel):
    """单条未命中提及的 LLM 归派候选。"""

    model_config = ConfigDict(extra="forbid")

    mention: str = Field(min_length=1, max_length=60)
    is_skill: bool
    capability_id: str | None = Field(default=None, pattern=r"^cap_\d{2}$")
    confidence: float = Field(ge=0.0, le=1.0)
    reason: str = Field(default="", max_length=30)


class ResumeAliasList(BaseModel):
    """对应 prompts/resume-alias.yml 的 output_schema。"""

    model_config = ConfigDict(extra="forbid")

    results: list[ResumeAliasResult] = Field(default_factory=list, max_length=40)


async def llm_resolve_unmatched(mentions: list[str], context: str) -> dict[str, ResumeAliasResult]:
    """LLM 归一层：词典未命中的提及批量归派，返回 mention -> 候选。

    只产出候选（capability_id + 置信度 + 理由），置信度 >= 0.6 才被采用；
    打分与岗位侧归一不受影响（岗位/日报侧仍纯词典，回测可复现）。
    """
    from ..infra.llm.promptspec import load_prompt
    from ..infra.llm.provider import build_provider
    from ..infra.llm.settings import LLMSettings
    from ..skills.resolver import load_resolver

    if not mentions:
        return {}
    spec = load_prompt("resume-alias")
    provider = build_provider(LLMSettings())
    resolver = load_resolver()
    caps = "\n".join(
        f"{e.capability_id} {e.name}（技能点: {'、'.join(p[0] for p in e.points) or '无'}）"
        for e in resolver.entries
    )
    items = "\n".join(f"- {m}（上下文: {context[:120]}）" for m in mentions)
    message = (
        f"未命中的技能提及:\n{items}\n\n能力域清单:\n{caps}\n\n任务: {spec.task}。只输出 JSON。"
    )
    last_err = "unknown"
    for attempt in range(2):
        user = message if attempt == 0 else f"{message}\n\n上次失败: {last_err}\n重新输出 JSON。"
        try:
            raw = await provider.complete(
                system=spec.system,
                user=user,
                max_tokens=int(spec.limits.get("max_tokens", 1500)),
                timeout=float(spec.limits.get("timeout_seconds", 120)),
            )
            t = raw.strip()
            if t.startswith("```"):
                t = t.split("\n", 1)[1]
                if t.rstrip().endswith("```"):
                    t = t.rstrip()[:-3]
            data = json.loads(t)
            parsed = ResumeAliasList.model_validate(data)
            return {r.mention: r for r in parsed.results}
        except Exception as exc:  # noqa: BLE001
            last_err = f"{type(exc).__name__}: {exc}"[:150]
    return {}


import json  # noqa: E402


def _parse_llm(raw: str) -> dict:
    """LLM 原文 -> ResumeRawExtraction dump；超限截断而不是整单拒绝。"""
    t = raw.strip()
    if t.startswith("```"):
        t = t.split("\n", 1)[1]
        if t.rstrip().endswith("```"):
            t = t.rstrip()[:-3]
    data = json.loads(t)
    if not isinstance(data, dict) or "skills" not in data:
        raise ValueError("输出缺少 skills")
    # ponytail: prompt 限不住数量，超限时截断而不是整单拒绝（回归实测 long 简历 >40 条触发）
    data["skills"] = data["skills"][:40]
    data["projects"] = data.get("projects", [])[:8]
    proficiency_map = {"熟悉": "中级", "熟练": "中级", "精通": "高级", "了解": "初级"}
    for skill in data["skills"]:
        if skill.get("proficiency") in proficiency_map:
            skill["proficiency"] = proficiency_map[skill["proficiency"]]
    for project in data["projects"]:
        project["skill_mentions"] = project.get("skill_mentions", [])[:15]
    data["work_experiences"] = data.get("work_experiences", [])[:10]
    data["education_history"] = data.get("education_history", [])[:8]
    data["certificates"] = data.get("certificates", [])[:12]
    return ResumeRawExtraction.model_validate(data).model_dump()


async def extract_resume(text: str, candidate_id: str = "demo") -> dict:
    """简历全文 -> LLM 结构化抽取（resume-parse.yml，重试 2 次）。

    被 candmatch/reseval/API 共用的域入口。
    """
    from ..infra.llm.promptspec import load_prompt
    from ..infra.llm.provider import build_provider
    from ..infra.llm.settings import LLMSettings

    spec = load_prompt("resume-parse")
    provider = build_provider(LLMSettings())
    message = f"candidate_id: {candidate_id}\n\n简历全文:\n{text[:8000]}"
    last_err = "unknown"
    for attempt in range(3):
        user = message if attempt == 0 else f"{message}\n\n上次失败: {last_err}\n重新输出 JSON。"
        try:
            raw = await provider.complete(
                system=spec.system,
                user=user,
                max_tokens=int(spec.limits.get("max_tokens", 2000)),
                timeout=float(spec.limits.get("timeout_seconds", 120)),
            )
            return _parse_llm(raw)
        except Exception as exc:  # noqa: BLE001
            last_err = f"{type(exc).__name__}: {exc}"[:200]
    raise RuntimeError(f"简历抽取失败: {last_err}")
